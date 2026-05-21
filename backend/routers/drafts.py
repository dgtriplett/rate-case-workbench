"""Download / upload / AI-revise endpoints for testimony + response drafts.

Lets a user round-trip drafts to Word (.docx) or Markdown / plain text — edit
externally — re-upload to overwrite the in-app draft, AND lets the user iterate
in-app with natural-language AI revision and one-click auto-apply of all
checklist recommendations.
"""
from __future__ import annotations

import io
import logging
import uuid
from typing import Optional

from fastapi import APIRouter, File, Form, HTTPException, Query, UploadFile
from fastapi.responses import Response
from pydantic import BaseModel
from sqlalchemy import select

from .checklist import EvalRequest, _load_items, _resolve_text, evaluate
from ..deps import CurrentUser, DBSession
from ..models import (
    DataRequest,
    Response as ResponseModel,
    ResponseStatus,
    Testimony,
)
from ..services.audit import log_event
from ..services.llm import chat as llm_chat

log = logging.getLogger(__name__)
router = APIRouter(prefix="/drafts", tags=["drafts"])


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


async def _load_target(session, kind: str, target_id: uuid.UUID):
    if kind == "testimony":
        res = await session.execute(select(Testimony).where(Testimony.id == target_id))
        t = res.scalar_one_or_none()
        if not t:
            raise HTTPException(404, "testimony not found")
        return t
    if kind == "response":
        res = await session.execute(select(ResponseModel).where(ResponseModel.id == target_id))
        r = res.scalar_one_or_none()
        if not r:
            raise HTTPException(404, "response not found")
        return r
    raise HTTPException(400, "kind must be 'testimony' or 'response'")


def _current_text(obj) -> str:
    return getattr(obj, "final_text", None) or getattr(obj, "draft_text", None) or ""


def _set_draft(obj, text_value: str) -> None:
    obj.draft_text = text_value
    # Don't auto-overwrite final_text — we only treat the draft as the working copy.
    if hasattr(obj, "status") and obj.status == ResponseStatus.filed:
        # Don't allow editing filed material via this path.
        raise HTTPException(409, "cannot edit a filed item")


def _title(obj) -> str:
    return getattr(obj, "title", None) or "Draft"


def _docx_bytes(title: str, body: str) -> bytes:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    doc.add_heading(title, level=1)
    for para in body.split("\n\n"):
        para = para.strip()
        if not para:
            continue
        doc.add_paragraph(para)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _docx_to_text(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())


# ---------------------------------------------------------------------------
# Download
# ---------------------------------------------------------------------------


@router.get("/{kind}/{target_id}/download")
async def download_draft(
    kind: str,
    target_id: uuid.UUID,
    session: DBSession,
    _: CurrentUser,
    fmt: str = Query("docx", pattern="^(docx|md|txt)$"),
) -> Response:
    obj = await _load_target(session, kind, target_id)
    text = _current_text(obj)
    title = _title(obj)
    safe_title = title.encode("ascii", errors="replace").decode("ascii")

    if fmt == "docx":
        data = _docx_bytes(title, text or "(empty draft)")
        media = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ext = "docx"
    elif fmt == "md":
        data = (f"# {title}\n\n{text}".encode("utf-8"))
        media = "text/markdown; charset=utf-8"
        ext = "md"
    else:
        data = (text or "").encode("utf-8")
        media = "text/plain; charset=utf-8"
        ext = "txt"

    file_safe = "".join(c if c.isalnum() else "-" for c in safe_title).strip("-")[:80] or "draft"
    return Response(
        content=data,
        media_type=media,
        headers={"Content-Disposition": f'attachment; filename="{file_safe}.{ext}"'},
    )


# ---------------------------------------------------------------------------
# Upload (replace draft)
# ---------------------------------------------------------------------------


class UploadResult(BaseModel):
    id: uuid.UUID
    bytes_received: int
    chars_imported: int


@router.post("/{kind}/{target_id}/upload", response_model=UploadResult)
async def upload_draft(
    kind: str,
    target_id: uuid.UUID,
    session: DBSession,
    user: CurrentUser,
    file: UploadFile = File(...),
) -> UploadResult:
    obj = await _load_target(session, kind, target_id)
    data = await file.read()
    if not data:
        raise HTTPException(400, "empty file")
    name = (file.filename or "").lower()
    if name.endswith(".docx"):
        text = _docx_to_text(data)
    elif name.endswith((".md", ".txt", ".markdown")):
        text = data.decode("utf-8", errors="replace")
    else:
        # Try utf-8 first, otherwise fall back to docx
        try:
            text = data.decode("utf-8")
        except UnicodeDecodeError:
            try:
                text = _docx_to_text(data)
            except Exception:
                raise HTTPException(400, "unsupported file type — upload .docx, .md, or .txt")

    if not text.strip():
        raise HTTPException(400, "uploaded file appears empty after parsing")
    _set_draft(obj, text)
    await session.flush()
    await log_event(
        session, actor=user, verb=f"{kind}.uploaded_revision",
        target_kind=kind, target_id=obj.id,
        case_id=getattr(obj, "case_id", None),
        payload={"filename": file.filename, "bytes": len(data), "chars": len(text)},
    )
    return UploadResult(id=obj.id, bytes_received=len(data), chars_imported=len(text))


# ---------------------------------------------------------------------------
# AI revise — natural language feedback loop
# ---------------------------------------------------------------------------


class ReviseIn(BaseModel):
    instruction: str
    additional_context: Optional[str] = None


class ReviseOut(BaseModel):
    id: uuid.UUID
    new_text: str
    summary: str


@router.post("/{kind}/{target_id}/revise", response_model=ReviseOut)
async def ai_revise(
    kind: str,
    target_id: uuid.UUID,
    body: ReviseIn,
    session: DBSession,
    user: CurrentUser,
) -> ReviseOut:
    obj = await _load_target(session, kind, target_id)
    current = _current_text(obj)
    title = _title(obj)
    prompt = (
        f"You are revising a {kind} draft titled '{title}'. Apply the user's "
        "instruction below to the draft. Return ONLY the full revised text — "
        "no preamble, no markdown fences, no commentary.\n\n"
        f"USER INSTRUCTION:\n{body.instruction}\n\n"
        + (f"ADDITIONAL CONTEXT:\n{body.additional_context}\n\n" if body.additional_context else "")
        + f"CURRENT DRAFT:\n{current}\n"
    )
    try:
        revised = await llm_chat(
            [{"role": "user", "content": prompt}], max_tokens=4000, temperature=0.3
        )
    except Exception as e:
        log.exception("revise LLM call failed")
        raise HTTPException(500, f"AI revision failed: {e}")

    revised = revised.strip()
    if revised.startswith("```"):
        revised = revised.split("\n", 1)[-1].rstrip("`").rstrip()

    _set_draft(obj, revised)
    await session.flush()

    # Produce a one-line summary for the audit log
    summary_prompt = (
        "In one short sentence, describe the change between this OLD and NEW draft. "
        "Be specific.\n\nOLD:\n" + (current[-2000:] or "(empty)") +
        "\n\nNEW:\n" + revised[-2000:]
    )
    try:
        summary = (await llm_chat(
            [{"role": "user", "content": summary_prompt}], max_tokens=80, temperature=0.0
        )).strip()
    except Exception:
        summary = body.instruction[:160]

    await log_event(
        session, actor=user, verb=f"{kind}.ai_revised",
        target_kind=kind, target_id=obj.id,
        case_id=getattr(obj, "case_id", None),
        payload={"instruction": body.instruction, "summary": summary,
                 "chars_before": len(current), "chars_after": len(revised)},
    )
    return ReviseOut(id=obj.id, new_text=revised, summary=summary)


# ---------------------------------------------------------------------------
# Auto-fix — run checklist + apply all suggested addenda in one shot
# ---------------------------------------------------------------------------


class AutoFixOut(BaseModel):
    id: uuid.UUID
    new_text: str
    applied_items: list[dict]
    skipped_items: list[dict]


@router.post("/{kind}/{target_id}/auto-fix", response_model=AutoFixOut)
async def auto_fix(
    kind: str,
    target_id: uuid.UUID,
    session: DBSession,
    user: CurrentUser,
) -> AutoFixOut:
    if kind not in ("testimony", "response"):
        raise HTTPException(400, "kind must be 'testimony' or 'response'")
    obj = await _load_target(session, kind, target_id)
    current = _current_text(obj)

    # Run checklist evaluation against the current draft
    eval_result = await evaluate(
        EvalRequest(kind=kind, target_id=target_id, case_id=getattr(obj, "case_id", None)),
        session, user,
    )

    failing = [
        i for i in eval_result.items
        if i.verdict in ("fail", "needs_attention") and i.suggested_addendum
    ]
    skipped = [
        {"id": i.id, "title": i.title, "verdict": i.verdict, "reason": i.rationale}
        for i in eval_result.items
        if i.verdict in ("fail", "needs_attention") and not i.suggested_addendum
    ]

    if not failing:
        return AutoFixOut(
            id=obj.id, new_text=current,
            applied_items=[], skipped_items=skipped,
        )

    # Apply the suggested addenda by asking the LLM to integrate them coherently
    addenda_block = "\n\n".join(
        f"[{i.id} — {i.title}]\n{i.suggested_addendum}" for i in failing
    )
    integrate_prompt = (
        "You are integrating reviewer-suggested additions into a draft. The "
        "current draft is below, followed by suggested additions (each one "
        "addresses a specific best-practices checklist item that is currently "
        "failing or needs attention). Integrate ALL of them into the draft, in "
        "the right place, with smooth transitions. Preserve the existing "
        "structure and the author's voice. Return ONLY the full revised draft "
        "— no commentary, no markdown fences.\n\n"
        f"CURRENT DRAFT:\n{current}\n\n"
        f"SUGGESTED ADDITIONS:\n{addenda_block}"
    )
    try:
        new_text = await llm_chat(
            [{"role": "user", "content": integrate_prompt}],
            max_tokens=4500, temperature=0.2,
        )
    except Exception as e:
        log.exception("auto-fix LLM call failed")
        raise HTTPException(500, f"auto-fix failed: {e}")

    new_text = new_text.strip()
    if new_text.startswith("```"):
        new_text = new_text.split("\n", 1)[-1].rstrip("`").rstrip()

    _set_draft(obj, new_text)
    await session.flush()
    applied = [{"id": i.id, "title": i.title, "verdict": i.verdict} for i in failing]
    await log_event(
        session, actor=user, verb=f"{kind}.auto_fix",
        target_kind=kind, target_id=obj.id,
        case_id=getattr(obj, "case_id", None),
        payload={"applied": [a["id"] for a in applied], "skipped": [s["id"] for s in skipped]},
    )
    return AutoFixOut(
        id=obj.id, new_text=new_text,
        applied_items=applied, skipped_items=skipped,
    )
